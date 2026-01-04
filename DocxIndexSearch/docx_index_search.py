"""
DOCX Index Search - A simple tool to index and search DOCX files
Single-file Windows application with SQLite database backend
"""

import sys
import os
import sqlite3
import subprocess
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from pathlib import Path
import json
from datetime import datetime

# Dependency check flag
DEPENDENCIES_OK = True
MISSING_DEPS = []

# Try importing required dependencies
try:
    from docx import Document
except ImportError:
    DEPENDENCIES_OK = False
    MISSING_DEPS.append("python-docx")

try:
    import PyPDF2
except ImportError:
    DEPENDENCIES_OK = False
    MISSING_DEPS.append("PyPDF2")


class DependencyChecker:
    """Check and optionally install missing dependencies"""

    @staticmethod
    def check_dependencies():
        """Check if all required dependencies are available"""
        return DEPENDENCIES_OK, MISSING_DEPS

    @staticmethod
    def install_dependencies(parent=None):
        """Attempt to install missing dependencies"""
        if not MISSING_DEPS:
            return True

        deps_str = ", ".join(MISSING_DEPS)
        message = f"Missing dependencies: {deps_str}\n\nWould you like to install them now?"

        if parent:
            result = messagebox.askyesno("Install Dependencies", message)
        else:
            print(message)
            result = input("Install? (y/n): ").lower() == 'y'

        if result:
            try:
                for dep in MISSING_DEPS:
                    print(f"Installing {dep}...")
                    subprocess.check_call([sys.executable, "-m", "pip", "install", dep])

                messagebox.showinfo("Success",
                    "Dependencies installed successfully!\nPlease restart the application.")
                return False  # Return False to indicate restart needed
            except Exception as e:
                error_msg = f"Failed to install dependencies: {str(e)}\n\n" \
                           f"Please install manually using:\npip install {' '.join(MISSING_DEPS)}"
                if parent:
                    messagebox.showerror("Installation Failed", error_msg)
                else:
                    print(error_msg)
                return False
        return False


class DocxDatabase:
    """SQLite database manager for DOCX file indexing"""

    def __init__(self, db_path="docx_index.db"):
        self.db_path = db_path
        self.conn = None
        self.init_database()

    def init_database(self):
        """Initialize database connection and create tables if needed"""
        self.conn = sqlite3.connect(self.db_path)
        self.conn.row_factory = sqlite3.Row
        cursor = self.conn.cursor()

        # Create documents table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT NOT NULL,
                filepath TEXT NOT NULL UNIQUE,
                file_size INTEGER,
                modified_date TEXT,
                indexed_date TEXT,
                content TEXT,
                file_type TEXT DEFAULT 'docx'
            )
        """)

        # Add file_type column if it doesn't exist (for existing databases)
        try:
            cursor.execute("ALTER TABLE documents ADD COLUMN file_type TEXT DEFAULT 'docx'")
        except sqlite3.OperationalError:
            pass  # Column already exists

        # Create settings table
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS settings (
                key TEXT PRIMARY KEY,
                value TEXT
            )
        """)

        # Create indexed_directories table for multiple directory support
        cursor.execute("""
            CREATE TABLE IF NOT EXISTS indexed_directories (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                path TEXT NOT NULL UNIQUE,
                indexed_date TEXT,
                enabled INTEGER DEFAULT 1
            )
        """)

        self.conn.commit()

    def get_setting(self, key, default=None):
        """Get a setting value"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = ?", (key,))
        row = cursor.fetchone()
        return row[0] if row else default

    def set_setting(self, key, value):
        """Set a setting value"""
        cursor = self.conn.cursor()
        cursor.execute("""
            INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)
        """, (key, value))
        self.conn.commit()

    def add_document(self, filename, filepath, content, file_size, modified_date, file_type='docx'):
        """Add or update a document in the index"""
        cursor = self.conn.cursor()
        indexed_date = datetime.now().isoformat()

        cursor.execute("""
            INSERT OR REPLACE INTO documents
            (filename, filepath, content, file_size, modified_date, indexed_date, file_type)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (filename, filepath, content, file_size, modified_date, indexed_date, file_type))

        self.conn.commit()

    def search_documents(self, query, file_types=None):
        """Search documents by content or filename
        Supports:
        - Exact phrase search: "The Abrahami Prophecy" (with quotes)
        - Multi-word search: The Abrahami Prophecy (all words must be present)
        - AND operator: James AND John (explicit AND)
        """
        cursor = self.conn.cursor()
        import re

        # Build file type filter
        type_filter = ""
        type_params = []
        if file_types:
            placeholders = ','.join('?' * len(file_types))
            type_filter = f" AND file_type IN ({placeholders})"
            type_params = list(file_types)

        # Check if query is wrapped in quotes for exact phrase search
        quote_pattern = re.compile(r'^["\'](.+)["\']$')
        quote_match = quote_pattern.match(query)

        if quote_match:
            phrase = quote_match.group(1)
            search_pattern = f"%{phrase}%"
            sql = f"""
                SELECT * FROM documents
                WHERE (content LIKE ? OR filename LIKE ?){type_filter}
                ORDER BY modified_date DESC
            """
            cursor.execute(sql, [search_pattern, search_pattern] + type_params)
        else:
            and_pattern = re.compile(r'\s+AND\s+', re.IGNORECASE)

            if and_pattern.search(query):
                terms = and_pattern.split(query)
                terms = [term.strip() for term in terms if term.strip()]

                if not terms:
                    return []

                conditions = []
                params = []

                for term in terms:
                    search_pattern = f"%{term}%"
                    conditions.append("(content LIKE ? OR filename LIKE ?)")
                    params.extend([search_pattern, search_pattern])

                sql = f"""
                    SELECT * FROM documents
                    WHERE {' AND '.join(conditions)}{type_filter}
                    ORDER BY modified_date DESC
                """
                cursor.execute(sql, params + type_params)
            else:
                words = query.split()

                if len(words) > 1:
                    # For multiple words, ALL words must appear in content (not split between filename and content)
                    content_conditions = []
                    filename_conditions = []
                    params = []

                    for word in words:
                        search_pattern = f"%{word}%"
                        content_conditions.append("content LIKE ?")
                        filename_conditions.append("filename LIKE ?")
                        params.append(search_pattern)  # for content

                    # Duplicate params for filename search
                    for word in words:
                        search_pattern = f"%{word}%"
                        params.append(search_pattern)  # for filename

                    # Match if ALL words are in content OR ALL words are in filename
                    sql = f"""
                        SELECT * FROM documents
                        WHERE (({' AND '.join(content_conditions)}) OR ({' AND '.join(filename_conditions)})){type_filter}
                        ORDER BY modified_date DESC
                    """
                    cursor.execute(sql, params + type_params)
                else:
                    search_pattern = f"%{query}%"
                    sql = f"""
                        SELECT * FROM documents
                        WHERE (content LIKE ? OR filename LIKE ?){type_filter}
                        ORDER BY modified_date DESC
                    """
                    cursor.execute(sql, [search_pattern, search_pattern] + type_params)

        return cursor.fetchall()

    def search_by_filename(self, query, file_types=None):
        """Search documents by filename only
        Supports AND/OR operators: David AND Saul, David OR Saul
        """
        cursor = self.conn.cursor()
        import re

        # Check for OR operator (case-insensitive)
        or_pattern = re.compile(r'\s+OR\s+', re.IGNORECASE)
        # Check for AND operator (case-insensitive)
        and_pattern = re.compile(r'\s+AND\s+', re.IGNORECASE)

        # Build file type filter
        type_filter = ""
        type_params = []
        if file_types:
            placeholders = ','.join('?' * len(file_types))
            type_filter = f" AND file_type IN ({placeholders})"
            type_params = list(file_types)

        if or_pattern.search(query):
            terms = or_pattern.split(query)
            terms = [term.strip() for term in terms if term.strip()]
            if not terms:
                return []
            conditions = []
            params = []
            for term in terms:
                conditions.append("filename LIKE ?")
                params.append(f"%{term}%")
            sql = f"""
                SELECT * FROM documents
                WHERE ({' OR '.join(conditions)}){type_filter}
                ORDER BY modified_date DESC
            """
            cursor.execute(sql, params + type_params)
        elif and_pattern.search(query):
            terms = and_pattern.split(query)
            terms = [term.strip() for term in terms if term.strip()]
            if not terms:
                return []
            conditions = []
            params = []
            for term in terms:
                conditions.append("filename LIKE ?")
                params.append(f"%{term}%")
            sql = f"""
                SELECT * FROM documents
                WHERE ({' AND '.join(conditions)}){type_filter}
                ORDER BY modified_date DESC
            """
            cursor.execute(sql, params + type_params)
        else:
            search_pattern = f"%{query}%"
            sql = f"""
                SELECT * FROM documents
                WHERE filename LIKE ?{type_filter}
                ORDER BY modified_date DESC
            """
            cursor.execute(sql, [search_pattern] + type_params)
        return cursor.fetchall()

    def search_content_in_files(self, content_query, filename_filter, file_types=None):
        """Search content only within files matching the filename filter
        Supports AND/OR operators in both filename_filter and content_query
        """
        cursor = self.conn.cursor()
        import re

        or_pattern = re.compile(r'\s+OR\s+', re.IGNORECASE)
        and_pattern = re.compile(r'\s+AND\s+', re.IGNORECASE)

        # Build filename conditions
        filename_conditions = []
        filename_params = []
        if or_pattern.search(filename_filter):
            terms = or_pattern.split(filename_filter)
            terms = [t.strip() for t in terms if t.strip()]
            for term in terms:
                filename_conditions.append("filename LIKE ?")
                filename_params.append(f"%{term}%")
            filename_sql = f"({' OR '.join(filename_conditions)})"
        elif and_pattern.search(filename_filter):
            terms = and_pattern.split(filename_filter)
            terms = [t.strip() for t in terms if t.strip()]
            for term in terms:
                filename_conditions.append("filename LIKE ?")
                filename_params.append(f"%{term}%")
            filename_sql = f"({' AND '.join(filename_conditions)})"
        else:
            filename_sql = "filename LIKE ?"
            filename_params.append(f"%{filename_filter}%")

        # Build content conditions
        content_conditions = []
        content_params = []
        if or_pattern.search(content_query):
            terms = or_pattern.split(content_query)
            terms = [t.strip() for t in terms if t.strip()]
            for term in terms:
                content_conditions.append("content LIKE ?")
                content_params.append(f"%{term}%")
            content_sql = f"({' OR '.join(content_conditions)})"
        elif and_pattern.search(content_query):
            terms = and_pattern.split(content_query)
            terms = [t.strip() for t in terms if t.strip()]
            for term in terms:
                content_conditions.append("content LIKE ?")
                content_params.append(f"%{term}%")
            content_sql = f"({' AND '.join(content_conditions)})"
        else:
            content_sql = "content LIKE ?"
            content_params.append(f"%{content_query}%")

        # Build file type filter
        type_filter = ""
        type_params = []
        if file_types:
            placeholders = ','.join('?' * len(file_types))
            type_filter = f" AND file_type IN ({placeholders})"
            type_params = list(file_types)

        sql = f"""
            SELECT * FROM documents
            WHERE {filename_sql} AND {content_sql}{type_filter}
            ORDER BY modified_date DESC
        """
        cursor.execute(sql, filename_params + content_params + type_params)
        return cursor.fetchall()

    def get_all_documents(self, file_types=None):
        """Get all indexed documents, optionally filtered by file type"""
        cursor = self.conn.cursor()
        if file_types:
            placeholders = ','.join('?' * len(file_types))
            cursor.execute(f"SELECT * FROM documents WHERE file_type IN ({placeholders}) ORDER BY modified_date DESC", file_types)
        else:
            cursor.execute("SELECT * FROM documents ORDER BY modified_date DESC")
        return cursor.fetchall()

    def get_document_count(self):
        """Get total number of indexed documents"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM documents")
        return cursor.fetchone()[0]

    def clear_index(self):
        """Clear all documents from the index"""
        cursor = self.conn.cursor()
        cursor.execute("DELETE FROM documents")
        self.conn.commit()

    def add_directory(self, path):
        """Add a directory to the indexed directories list"""
        cursor = self.conn.cursor()
        try:
            cursor.execute("""
                INSERT INTO indexed_directories (path, indexed_date, enabled)
                VALUES (?, ?, 1)
            """, (path, datetime.now().isoformat()))
            self.conn.commit()
            return True
        except sqlite3.IntegrityError:
            # Directory already exists
            return False

    def remove_directory(self, path):
        """Remove a directory from the indexed directories list"""
        cursor = self.conn.cursor()
        cursor.execute("DELETE FROM indexed_directories WHERE path = ?", (path,))
        self.conn.commit()

    def get_directories(self):
        """Get all indexed directories"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM indexed_directories ORDER BY path")
        return cursor.fetchall()

    def get_enabled_directories(self):
        """Get only enabled directories"""
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM indexed_directories WHERE enabled = 1 ORDER BY path")
        return cursor.fetchall()

    def toggle_directory(self, path, enabled):
        """Enable or disable a directory"""
        cursor = self.conn.cursor()
        cursor.execute("""
            UPDATE indexed_directories SET enabled = ? WHERE path = ?
        """, (1 if enabled else 0, path))
        self.conn.commit()

    def close(self):
        """Close database connection"""
        if self.conn:
            self.conn.close()


class DocxIndexer:
    """DOCX and PDF file indexer"""

    @staticmethod
    def extract_text_from_docx(filepath):
        """Extract text content from a DOCX file"""
        try:
            doc = Document(filepath)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts)
        except Exception as e:
            print(f"Error extracting text from {filepath}: {e}")
            return ""

    @staticmethod
    def extract_text_from_pdf(filepath):
        """Extract text content from a PDF file"""
        try:
            text_parts = []
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text and text.strip():
                        text_parts.append(text)
            return "\n".join(text_parts)
        except Exception as e:
            print(f"Error extracting text from {filepath}: {e}")
            return ""

    @staticmethod
    def index_directory(directory, db, progress_callback=None):
        """Index all DOCX files in a directory"""
        docx_files = list(Path(directory).glob("**/*.docx"))
        # Filter out temporary files (starting with ~$)
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]

        total = len(docx_files)
        indexed = 0
        errors = []

        for i, filepath in enumerate(docx_files):
            try:
                content = DocxIndexer.extract_text_from_docx(str(filepath))
                stats = filepath.stat()

                db.add_document(
                    filename=filepath.name,
                    filepath=str(filepath.absolute()),
                    content=content,
                    file_size=stats.st_size,
                    modified_date=datetime.fromtimestamp(stats.st_mtime).isoformat()
                )

                indexed += 1

                if progress_callback:
                    progress_callback(i + 1, total, filepath.name)

            except Exception as e:
                errors.append(f"{filepath.name}: {str(e)}")
                print(f"Error indexing {filepath}: {e}")

        return indexed, errors

    @staticmethod
    def index_directories(directories, db, progress_callback=None):
        """Index all DOCX and PDF files from multiple directories"""
        total_indexed = 0
        all_errors = []

        # Collect all docx and pdf files from all directories
        all_files = []
        for directory in directories:
            # Get DOCX files
            docx_files = list(Path(directory).glob("**/*.docx"))
            docx_files = [f for f in docx_files if not f.name.startswith("~$")]
            all_files.extend([(f, 'docx') for f in docx_files])

            # Get PDF files
            pdf_files = list(Path(directory).glob("**/*.pdf"))
            pdf_files = [f for f in pdf_files if not f.name.startswith("~$")]
            all_files.extend([(f, 'pdf') for f in pdf_files])

        total = len(all_files)
        current = 0

        # Index all files
        for filepath, file_type in all_files:
            try:
                if file_type == 'docx':
                    content = DocxIndexer.extract_text_from_docx(str(filepath))
                else:
                    content = DocxIndexer.extract_text_from_pdf(str(filepath))

                stats = filepath.stat()

                db.add_document(
                    filename=filepath.name,
                    filepath=str(filepath.absolute()),
                    content=content,
                    file_size=stats.st_size,
                    modified_date=datetime.fromtimestamp(stats.st_mtime).isoformat(),
                    file_type=file_type
                )

                total_indexed += 1
                current += 1

                if progress_callback:
                    progress_callback(current, total, filepath.name)

            except Exception as e:
                all_errors.append(f"{filepath.name}: {str(e)}")
                print(f"Error indexing {filepath}: {e}")
                current += 1

        return total_indexed, all_errors


class DocxSearchApp:
    """Main application window"""

    def __init__(self, root):
        self.root = root
        self.root.title("DOCX Index Search")
        self.root.geometry("900x600")

        # Initialize database
        self.db = DocxDatabase()

        # Store current results for sorting
        self.current_results = []
        self.current_title = "All Documents"
        self.sort_column = None
        self.sort_reverse = False

        # Create UI first
        self.create_widgets()

        # Restore window settings
        self.restore_window_settings()

        # Check if this is first run (after UI is created)
        self.check_first_run()

        # Load initial data
        self.load_all_documents()

    def check_first_run(self):
        """Check if database is empty and offer to index"""
        doc_count = self.db.get_document_count()

        if doc_count == 0:
            result = messagebox.askyesno(
                "First Run - Create Index",
                "No documents found in the index.\n\n"
                "Would you like to scan for DOCX files and create an index now?"
            )

            if result:
                self.index_documents()

    def create_widgets(self):
        """Create the main UI widgets"""
        # Filename search frame
        filename_frame = ttk.Frame(self.root, padding="10 10 10 5")
        filename_frame.pack(fill=tk.X)

        ttk.Label(filename_frame, text="Filename Search:").pack(side=tk.LEFT, padx=(0, 5))
        self.filename_search_var = tk.StringVar()
        filename_entry = ttk.Entry(filename_frame, textvariable=self.filename_search_var, width=50)
        filename_entry.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(filename_frame, text="Search", command=self.perform_filename_search).pack(side=tk.LEFT, padx=2)

        # Content search frame
        content_frame = ttk.Frame(self.root, padding="10 5 10 5")
        content_frame.pack(fill=tk.X)

        ttk.Label(content_frame, text="Content Search:").pack(side=tk.LEFT, padx=(0, 5))
        self.search_var = tk.StringVar()
        self.search_var.trace("w", self.on_search_change)
        search_entry = ttk.Entry(content_frame, textvariable=self.search_var, width=50)
        search_entry.pack(side=tk.LEFT, padx=(0, 10))
        ttk.Button(content_frame, text="Search", command=self.perform_search).pack(side=tk.LEFT, padx=2)

        # File type checkboxes and buttons frame
        buttons_frame = ttk.Frame(self.root, padding="10 5 10 5")
        buttons_frame.pack(fill=tk.X)

        # File type checkboxes
        ttk.Label(buttons_frame, text="File Types:").pack(side=tk.LEFT, padx=(0, 5))
        self.docx_var = tk.BooleanVar(value=True)
        self.pdf_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(buttons_frame, text="DOCX", variable=self.docx_var).pack(side=tk.LEFT, padx=2)
        ttk.Checkbutton(buttons_frame, text="PDF", variable=self.pdf_var).pack(side=tk.LEFT, padx=2)

        ttk.Separator(buttons_frame, orient='vertical').pack(side=tk.LEFT, padx=10, fill=tk.Y)

        ttk.Button(buttons_frame, text="Show All", command=self.load_all_documents).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Re-index", command=self.index_documents).pack(side=tk.LEFT, padx=2)
        ttk.Button(buttons_frame, text="Settings", command=self.show_settings).pack(side=tk.LEFT, padx=2)

        # Search hint
        hint_frame = ttk.Frame(self.root, padding="0 0 10 5")
        hint_frame.pack(fill=tk.X)
        hint_label = ttk.Label(hint_frame, text='Tip: Use quotes for exact phrases ("The Prophecy") or type words for all-words search',
                              font=('TkDefaultFont', 8), foreground='gray')
        hint_label.pack(side=tk.RIGHT)

        # Middle frame for results
        middle_frame = ttk.Frame(self.root, padding="10")
        middle_frame.pack(fill=tk.BOTH, expand=True)

        # Results label
        self.results_label = ttk.Label(middle_frame, text="All Documents (0)")
        self.results_label.pack(anchor=tk.W, pady=(0, 5))

        # Results treeview with scrollbar
        tree_frame = ttk.Frame(middle_frame)
        tree_frame.pack(fill=tk.BOTH, expand=True)

        # Scrollbars
        vsb = ttk.Scrollbar(tree_frame, orient="vertical")
        hsb = ttk.Scrollbar(tree_frame, orient="horizontal")

        # Treeview
        self.results_tree = ttk.Treeview(
            tree_frame,
            columns=("filename", "modified", "size", "path"),
            show="headings",
            yscrollcommand=vsb.set,
            xscrollcommand=hsb.set
        )

        vsb.config(command=self.results_tree.yview)
        hsb.config(command=self.results_tree.xview)

        # Configure columns with sorting
        self.results_tree.heading("filename", text="Filename",
                                  command=lambda: self.sort_by_column("filename"))
        self.results_tree.heading("modified", text="Modified",
                                  command=lambda: self.sort_by_column("modified"))
        self.results_tree.heading("size", text="Size",
                                  command=lambda: self.sort_by_column("size"))
        self.results_tree.heading("path", text="Path",
                                  command=lambda: self.sort_by_column("path"))

        self.results_tree.column("filename", width=250)
        self.results_tree.column("modified", width=150)
        self.results_tree.column("size", width=100)
        self.results_tree.column("path", width=350)

        # Pack scrollbars and treeview
        self.results_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        hsb.pack(side=tk.BOTTOM, fill=tk.X)

        # Bind double-click to open file
        self.results_tree.bind("<Double-1>", self.open_document)

        # Status bar
        self.status_var = tk.StringVar()
        self.status_var.set("Ready")
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X)

    def on_search_change(self, *args):
        """Called when search text changes"""
        # Auto-search as user types (with a small delay this would be better, but keeping it simple)
        pass

    def perform_filename_search(self):
        """Search only by filename"""
        query = self.filename_search_var.get().strip()

        if not query:
            self.load_all_documents()
            return

        file_types = self.get_selected_file_types()
        self.status_var.set(f"Searching filenames for: {query}")
        results = self.db.search_by_filename(query, file_types)
        self.display_results(results, f"Filename Search Results for '{query}'")
        self.status_var.set(f"Found {len(results)} document(s)")

    def perform_search(self):
        """Perform search based on search box content, filtered by filename if specified"""
        query = self.search_var.get().strip()
        filename_filter = self.filename_search_var.get().strip()
        file_types = self.get_selected_file_types()

        if not query:
            self.load_all_documents()
            return

        if filename_filter:
            self.status_var.set(f"Searching content for: {query} (in files matching '{filename_filter}')")
            results = self.db.search_content_in_files(query, filename_filter, file_types)
            self.display_results(results, f"Content '{query}' in files matching '{filename_filter}'")
        else:
            self.status_var.set(f"Searching for: {query}")
            results = self.db.search_documents(query, file_types)
            self.display_results(results, f"Search Results for '{query}'")
        self.status_var.set(f"Found {len(results)} document(s)")

    def get_selected_file_types(self):
        """Get list of selected file types from checkboxes"""
        types = []
        if hasattr(self, 'docx_var') and self.docx_var.get():
            types.append('docx')
        if hasattr(self, 'pdf_var') and self.pdf_var.get():
            types.append('pdf')
        return types if types else None

    def load_all_documents(self):
        """Load and display all documents"""
        # Safety check - ensure UI is created
        if not hasattr(self, 'search_var'):
            return

        self.search_var.set("")
        self.filename_search_var.set("")
        file_types = self.get_selected_file_types()
        results = self.db.get_all_documents(file_types)
        self.display_results(results, "All Documents")
        self.status_var.set(f"Showing {len(results)} document(s)")

    def display_results(self, results, title="Results"):
        """Display search results in the treeview"""
        # Safety check - ensure UI is created
        if not hasattr(self, 'results_tree'):
            return

        # Store current results for sorting
        self.current_results = list(results)
        self.current_title = title

        # Reset sort state when showing new results
        self.sort_column = None
        self.sort_reverse = False
        self.update_column_headings()

        # Clear existing items
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)

        # Update label
        self.results_label.config(text=f"{title} ({len(results)})")

        # Add results
        for row in results:
            # Format file size
            size_kb = row['file_size'] / 1024 if row['file_size'] else 0
            size_str = f"{size_kb:.1f} KB"

            # Format modified date
            try:
                modified = datetime.fromisoformat(row['modified_date'])
                modified_str = modified.strftime("%Y-%m-%d %H:%M")
            except:
                modified_str = row['modified_date']

            self.results_tree.insert("", tk.END, values=(
                row['filename'],
                modified_str,
                size_str,
                row['filepath']
            ))

    def sort_by_column(self, column):
        """Sort results by the specified column"""
        if not self.current_results:
            return

        # Toggle sort direction if clicking the same column
        if self.sort_column == column:
            self.sort_reverse = not self.sort_reverse
        else:
            self.sort_column = column
            self.sort_reverse = False

        # Define sorting key functions
        def get_sort_key(row):
            if column == "filename":
                return row['filename'].lower()
            elif column == "modified":
                return row['modified_date']
            elif column == "size":
                return row['file_size'] or 0
            elif column == "path":
                return row['filepath'].lower()
            return ""

        # Sort the results
        sorted_results = sorted(self.current_results, key=get_sort_key, reverse=self.sort_reverse)

        # Update column headings with sort indicators
        self.update_column_headings()

        # Display sorted results
        self._display_sorted_results(sorted_results)

    def update_column_headings(self):
        """Update column headings to show sort indicators"""
        # Unicode arrows for sort direction
        up_arrow = " ▲"
        down_arrow = " ▼"

        # Base column names
        columns = {
            "filename": "Filename",
            "modified": "Modified",
            "size": "Size",
            "path": "Path"
        }

        # Update each heading
        for col_id, col_name in columns.items():
            if col_id == self.sort_column:
                arrow = down_arrow if self.sort_reverse else up_arrow
                self.results_tree.heading(col_id, text=f"{col_name}{arrow}")
            else:
                self.results_tree.heading(col_id, text=col_name)

    def _display_sorted_results(self, results):
        """Display pre-sorted results without storing them again"""
        # Clear existing items
        for item in self.results_tree.get_children():
            self.results_tree.delete(item)

        # Add results
        for row in results:
            # Format file size
            size_kb = row['file_size'] / 1024 if row['file_size'] else 0
            size_str = f"{size_kb:.1f} KB"

            # Format modified date
            try:
                modified = datetime.fromisoformat(row['modified_date'])
                modified_str = modified.strftime("%Y-%m-%d %H:%M")
            except:
                modified_str = row['modified_date']

            self.results_tree.insert("", tk.END, values=(
                row['filename'],
                modified_str,
                size_str,
                row['filepath']
            ))

    def open_document(self, event):
        """Open selected document in MS Word"""
        selection = self.results_tree.selection()
        if not selection:
            return

        item = self.results_tree.item(selection[0])
        filepath = item['values'][3]  # Path is the 4th column

        if not os.path.exists(filepath):
            messagebox.showerror("File Not Found",
                f"The file no longer exists:\n{filepath}")
            return

        try:
            # Open file with default application (MS Word)
            os.startfile(filepath)
            self.status_var.set(f"Opened: {os.path.basename(filepath)}")
        except Exception as e:
            messagebox.showerror("Error Opening File",
                f"Could not open file:\n{str(e)}")

    def index_documents(self):
        """Index documents from configured directories"""
        # Get all enabled directories
        directories = self.db.get_enabled_directories()

        if not directories:
            # No directories configured, prompt to add one
            result = messagebox.askyesno(
                "No Directories Configured",
                "No directories are configured for indexing.\n\n"
                "Would you like to add a directory now?"
            )
            if result:
                self.show_settings()
            return

        # Show which directories will be indexed
        dir_list = "\n".join([d['path'] for d in directories])
        result = messagebox.askyesno(
            "Index Documents",
            f"Index DOCX files from the following directories:\n\n{dir_list}\n\n"
            "Click Yes to proceed\n"
            "Click No to cancel"
        )

        if not result:
            return

        # Create progress window
        progress_window = tk.Toplevel(self.root)
        progress_window.title("Indexing Documents")
        progress_window.geometry("400x150")
        progress_window.transient(self.root)
        progress_window.grab_set()

        ttk.Label(progress_window, text="Indexing DOCX files from multiple directories...").pack(pady=10)

        progress_var = tk.StringVar()
        progress_label = ttk.Label(progress_window, textvariable=progress_var)
        progress_label.pack(pady=5)

        progress_bar = ttk.Progressbar(progress_window, mode='determinate', length=300)
        progress_bar.pack(pady=10)

        def progress_callback(current, total, filename):
            progress_var.set(f"Processing {current}/{total}: {filename}")
            progress_bar['maximum'] = total
            progress_bar['value'] = current
            progress_window.update()

        # Perform indexing
        try:
            # Get list of directory paths
            dir_paths = [d['path'] for d in directories]

            indexed, errors = DocxIndexer.index_directories(
                dir_paths,
                self.db,
                progress_callback
            )

            progress_window.destroy()

            # Show results
            message = f"Successfully indexed {indexed} document(s) from {len(dir_paths)} director{'y' if len(dir_paths) == 1 else 'ies'}."
            if errors:
                message += f"\n\n{len(errors)} error(s) occurred:\n" + "\n".join(errors[:5])
                if len(errors) > 5:
                    message += f"\n... and {len(errors) - 5} more"

            messagebox.showinfo("Indexing Complete", message)

            # Refresh display
            self.load_all_documents()

        except Exception as e:
            progress_window.destroy()
            messagebox.showerror("Indexing Error", f"An error occurred:\n{str(e)}")

    def show_settings(self):
        """Show settings dialog with multiple directory management"""
        settings_window = tk.Toplevel(self.root)
        settings_window.title("Settings - Indexed Directories")
        settings_window.geometry("700x450")
        settings_window.transient(self.root)
        settings_window.grab_set()

        # Title
        ttk.Label(settings_window, text="Indexed Directories:", font=('TkDefaultFont', 10, 'bold')).pack(pady=(10, 5), padx=20, anchor=tk.W)

        # Frame for directory list and buttons
        list_frame = ttk.Frame(settings_window)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=5)

        # Listbox with scrollbar for directories
        scrollbar = ttk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        directories_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, height=10)
        directories_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=directories_listbox.yview)

        # Load current directories
        def refresh_directory_list():
            directories_listbox.delete(0, tk.END)
            dirs = self.db.get_directories()
            for dir_row in dirs:
                status = "✓" if dir_row['enabled'] else "✗"
                directories_listbox.insert(tk.END, f"{status} {dir_row['path']}")

        refresh_directory_list()

        # Buttons frame
        buttons_frame = ttk.Frame(settings_window)
        buttons_frame.pack(fill=tk.X, padx=20, pady=10)

        def add_directory():
            directory = filedialog.askdirectory(
                title="Select Directory to Add for Indexing"
            )
            if directory:
                if self.db.add_directory(directory):
                    refresh_directory_list()
                    messagebox.showinfo("Directory Added", f"Added directory:\n{directory}")
                else:
                    messagebox.showwarning("Already Exists", "This directory is already in the list.")

        def remove_directory():
            selection = directories_listbox.curselection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select a directory to remove.")
                return

            dirs = self.db.get_directories()
            if selection[0] < len(dirs):
                dir_path = dirs[selection[0]]['path']
                result = messagebox.askyesno(
                    "Remove Directory",
                    f"Remove this directory from the index list?\n\n{dir_path}\n\n"
                    "Note: This will not delete any indexed documents.\n"
                    "Use 'Clear Index' to remove indexed documents."
                )
                if result:
                    self.db.remove_directory(dir_path)
                    refresh_directory_list()

        def toggle_directory():
            selection = directories_listbox.curselection()
            if not selection:
                messagebox.showwarning("No Selection", "Please select a directory to enable/disable.")
                return

            dirs = self.db.get_directories()
            if selection[0] < len(dirs):
                dir_row = dirs[selection[0]]
                new_state = not dir_row['enabled']
                self.db.toggle_directory(dir_row['path'], new_state)
                refresh_directory_list()

        ttk.Button(buttons_frame, text="Add Directory", command=add_directory).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Remove Directory", command=remove_directory).pack(side=tk.LEFT, padx=5)
        ttk.Button(buttons_frame, text="Enable/Disable", command=toggle_directory).pack(side=tk.LEFT, padx=5)

        # Separator
        ttk.Separator(settings_window, orient='horizontal').pack(fill=tk.X, padx=20, pady=10)

        # Database info
        info_frame = ttk.Frame(settings_window)
        info_frame.pack(fill=tk.X, padx=20, pady=5)

        ttk.Label(info_frame, text=f"Database: {self.db.db_path}").pack(anchor=tk.W)
        ttk.Label(info_frame, text=f"Indexed Documents: {self.db.get_document_count()}").pack(anchor=tk.W)

        # Bottom buttons
        bottom_button_frame = ttk.Frame(settings_window)
        bottom_button_frame.pack(pady=10)

        def clear_index():
            result = messagebox.askyesno(
                "Clear Index",
                "Are you sure you want to clear the entire index?\n\n"
                "This will remove all indexed documents from the database.\n"
                "The actual files and directory list will not be deleted."
            )
            if result:
                self.db.clear_index()
                messagebox.showinfo("Index Cleared", "The index has been cleared.")
                self.load_all_documents()
                # Update document count
                for widget in info_frame.winfo_children():
                    widget.destroy()
                ttk.Label(info_frame, text=f"Database: {self.db.db_path}").pack(anchor=tk.W)
                ttk.Label(info_frame, text=f"Indexed Documents: {self.db.get_document_count()}").pack(anchor=tk.W)

        ttk.Button(bottom_button_frame, text="Clear Index", command=clear_index).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom_button_frame, text="Close", command=settings_window.destroy).pack(side=tk.LEFT, padx=5)

    def save_window_settings(self):
        """Save window geometry and settings"""
        geometry = self.root.geometry()
        self.db.set_setting('window_geometry', geometry)

    def restore_window_settings(self):
        """Restore window geometry and settings"""
        geometry = self.db.get_setting('window_geometry')
        if geometry:
            try:
                self.root.geometry(geometry)
            except:
                pass  # Use default if invalid

    def on_closing(self):
        """Handle application closing"""
        self.save_window_settings()
        self.db.close()
        self.root.destroy()


def main():
    """Main application entry point"""
    # Check dependencies first
    deps_ok, missing = DependencyChecker.check_dependencies()

    if not deps_ok:
        # Create a minimal root window for the dependency dialog
        root = tk.Tk()
        root.withdraw()  # Hide the main window

        deps_str = ", ".join(missing)
        messagebox.showwarning(
            "Missing Dependencies",
            f"The following required packages are missing:\n\n{deps_str}\n\n"
            "The application needs these to function properly."
        )

        installed = DependencyChecker.install_dependencies(root)

        if not installed:
            messagebox.showinfo(
                "Manual Installation Required",
                f"Please install the required packages manually:\n\n"
                f"pip install {' '.join(missing)}\n\n"
                "Then restart the application."
            )
            root.destroy()
            return

    # Create and run the main application
    root = tk.Tk()
    app = DocxSearchApp(root)
    root.protocol("WM_DELETE_WINDOW", app.on_closing)
    root.mainloop()


if __name__ == "__main__":
    main()
